from functools import partial
from operator import contains

from docx import Document
from docx.shared import RGBColor

document = Document("refer.docx")


def search_and_paint(names, years):
    found = False
    for paragraph in document.paragraphs:
        text = paragraph.text
        in_text = partial(contains, text)
        if all(map(in_text, names)) and any(map(in_text, years)):
            found = True
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return found


def save():
    document.save("new_refer.docx")


if __name__ == "__main__":
    for paragraph in document.paragraphs:
        print(paragraph.text)
